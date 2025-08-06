import os
import uuid
import io
from typing import List, Dict, Any
from datetime import datetime
import pypdf
from docx import Document
import email
from email import policy
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF file and extract text"""
        try:
            # Convert bytes to BytesIO object for pypdf
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "text": text,
                "pages": len(pdf_reader.pages),
                "filename": filename
            }
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX file and extract text"""
        try:
            # Convert bytes to BytesIO object for python-docx
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "text": text,
                "paragraphs": len(doc.paragraphs),
                "filename": filename
            }
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def process_email(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process email file and extract text"""
        try:
            msg = email.message_from_bytes(file_content, policy=policy.default)
            text = ""
            
            # Extract subject
            subject = msg.get('subject', '')
            text += f"Subject: {subject}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text += part.get_payload(decode=True).decode()
                        break
            else:
                text += msg.get_payload(decode=True).decode()
            
            return {
                "text": text,
                "subject": subject,
                "filename": filename
            }
        except Exception as e:
            raise Exception(f"Error processing email: {str(e)}")
    
    def chunk_text(self, text: str, document_id: str) -> List[LangchainDocument]:
        """Split text into chunks using LangChain"""
        try:
            chunks = self.text_splitter.split_text(text)
            documents = []
            
            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "document_id": document_id,
                        "chunk_id": i,
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            
            return documents
        except Exception as e:
            raise Exception(f"Error chunking text: {str(e)}")
    
    def process_document(self, file_content: bytes, filename: str, file_type: str) -> Dict[str, Any]:
        """Main method to process any document type"""
        document_id = str(uuid.uuid4())
        
        # Process based on file type
        if file_type.lower() == "pdf":
            result = self.process_pdf(file_content, filename)
        elif file_type.lower() in ["docx", "doc"]:
            result = self.process_docx(file_content, filename)
        elif file_type.lower() in ["eml", "email"]:
            result = self.process_email(file_content, filename)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
        
        # Chunk the text
        chunks = self.chunk_text(result["text"], document_id)
        
        return {
            "document_id": document_id,
            "filename": filename,
            "file_type": file_type,
            "text": result["text"],
            "chunks": chunks,
            "total_chunks": len(chunks),
            "upload_time": datetime.now()
        } 